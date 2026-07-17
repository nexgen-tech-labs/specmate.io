"""Router-level tests use a plain (sync) test function, not `async def` +
@pytest.mark.asyncio — TestClient drives the FastAPI app (and its own async DB session)
through its own event loop internally. Mixing that with an outer pytest-asyncio event
loop touching the same asyncpg engine causes 'another operation is in progress' /
MissingGreenlet errors, since asyncpg connections can't be shared across event loops.
Fixture setup/teardown therefore uses asyncio.run() — a fresh, self-contained loop per
call — rather than an async test function sharing a loop with TestClient."""

from __future__ import annotations

import asyncio
import io
from datetime import UTC, datetime

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import delete
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine

from reportlab.pdfgen import canvas

from app.core import db as db_module
from app.core.config import settings
from app.main import app
from app.models import Project, RawRequirement, Source, SourceDiff, Workspace
from app.routers.sources import get_blob_downloader
from app.services.parsing.docx_parser import DocxParseError
from app.services.storage.blob_client import BlobDownloadError
from tests.audit_cleanup import purge_audit_events


def _now() -> datetime:
    return datetime.now(UTC).replace(tzinfo=None)


def _dispose_app_engine() -> None:
    """TestClient runs the app (and app.core.db's module-level engine) through its own
    event loop per call. Once that loop closes, the engine's pooled asyncpg connections
    are unusable on the next call — dispose the pool so a fresh connection is opened."""
    asyncio.run(db_module.engine.dispose())


def _sample_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("1. Overview", level=1)
    document.add_paragraph("Some requirement text.")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _sample_pdf_bytes() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Some requirement text.")
    c.showPage()
    c.save()
    return buf.getvalue()


async def _create_source(kind: str = "DOCX", storage_key: str | None = "some/key.docx") -> dict[str, str]:
    engine = create_async_engine(settings.database_url)
    try:
        async with AsyncSession(engine) as session:
            workspace = Workspace(name="Parser Test WS", createdAt=_now(), updatedAt=_now())
            session.add(workspace)
            await session.flush()

            project = Project(
                workspaceId=workspace.id,
                name="Parser Test Project",
                createdAt=_now(),
                updatedAt=_now(),
            )
            session.add(project)
            await session.flush()

            source = Source(
                projectId=project.id,
                name="test.docx",
                kind=kind,
                storageKey=storage_key,
                createdAt=_now(),
                updatedAt=_now(),
            )
            session.add(source)
            await session.flush()

            ids = {
                "source_id": source.id,
                "project_id": project.id,
                "workspace_id": workspace.id,
            }
            await session.commit()
            return ids
    finally:
        await engine.dispose()


async def _create_new_version(
    project_id: str, previous_source_id: str, previous_version: int, storage_key: str
) -> str:
    engine = create_async_engine(settings.database_url)
    try:
        async with AsyncSession(engine) as session:
            source = Source(
                projectId=project_id,
                name="test.docx",
                kind="DOCX",
                storageKey=storage_key,
                version=previous_version + 1,
                previousVersionId=previous_source_id,
                createdAt=_now(),
                updatedAt=_now(),
            )
            session.add(source)
            await session.flush()
            new_id = source.id
            await session.commit()
            return new_id
    finally:
        await engine.dispose()


def _docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for p in paragraphs:
        document.add_paragraph(p)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


async def _fetch_source_status(source_id: str) -> str:
    engine = create_async_engine(settings.database_url)
    try:
        async with AsyncSession(engine) as session:
            source = await session.get(Source, source_id)
            assert source is not None
            status = source.status
            return str(status.value if hasattr(status, "value") else status)
    finally:
        await engine.dispose()


async def _count_raw_requirements(source_id: str) -> int:
    engine = create_async_engine(settings.database_url)
    try:
        async with AsyncSession(engine) as session:
            result = await session.execute(
                RawRequirement.__table__.select().where(RawRequirement.sourceId == source_id)
            )
            return len(result.fetchall())
    finally:
        await engine.dispose()


async def _cleanup(ids: dict[str, str], extra_source_ids: list[str] | None = None) -> None:
    engine = create_async_engine(settings.database_url)
    try:
        async with AsyncSession(engine) as session:
            all_source_ids = [ids["source_id"], *(extra_source_ids or [])]
            await session.execute(
                delete(SourceDiff).where(SourceDiff.sourceId.in_(all_source_ids))
            )
            await session.execute(
                delete(RawRequirement).where(RawRequirement.sourceId.in_(all_source_ids))
            )
            await session.execute(delete(Source).where(Source.id.in_(all_source_ids)))
            await session.execute(delete(Project).where(Project.id == ids["project_id"]))
            await purge_audit_events(session, ids["workspace_id"])
            await session.execute(delete(Workspace).where(Workspace.id == ids["workspace_id"]))
            await session.commit()
    finally:
        await engine.dispose()


def test_parse_creates_raw_requirements_and_marks_source_parsed() -> None:
    ids = asyncio.run(_create_source())

    async def fake_downloader(_storage_key: str) -> bytes:
        return _sample_docx_bytes()

    app.dependency_overrides[get_blob_downloader] = lambda: fake_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "PARSED"
    assert body["chunk_count"] == 1

    assert asyncio.run(_fetch_source_status(ids["source_id"])) == "PARSED"
    assert asyncio.run(_count_raw_requirements(ids["source_id"])) == 1

    asyncio.run(_cleanup(ids))


def test_parse_failure_marks_source_failed_and_returns_422_with_no_partial_rows() -> None:
    ids = asyncio.run(_create_source())

    async def failing_downloader(_storage_key: str) -> bytes:
        raise BlobDownloadError("simulated blob fetch failure")

    app.dependency_overrides[get_blob_downloader] = lambda: failing_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 422
    assert "simulated blob fetch failure" in response.json()["detail"]

    assert asyncio.run(_fetch_source_status(ids["source_id"])) == "FAILED"
    assert asyncio.run(_count_raw_requirements(ids["source_id"])) == 0

    asyncio.run(_cleanup(ids))


def test_docx_parse_error_also_marks_source_failed() -> None:
    ids = asyncio.run(_create_source())

    async def bad_content_downloader(_storage_key: str) -> bytes:
        raise DocxParseError("corrupt docx")

    app.dependency_overrides[get_blob_downloader] = lambda: bad_content_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 422
    assert asyncio.run(_fetch_source_status(ids["source_id"])) == "FAILED"

    asyncio.run(_cleanup(ids))


def test_pdf_parse_creates_raw_requirements_and_marks_source_parsed() -> None:
    ids = asyncio.run(_create_source(kind="PDF", storage_key="some/key.pdf"))

    async def fake_downloader(_storage_key: str) -> bytes:
        return _sample_pdf_bytes()

    app.dependency_overrides[get_blob_downloader] = lambda: fake_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "PARSED"
    assert body["chunk_count"] == 1

    assert asyncio.run(_fetch_source_status(ids["source_id"])) == "PARSED"
    assert asyncio.run(_count_raw_requirements(ids["source_id"])) == 1

    asyncio.run(_cleanup(ids))


def test_unsupported_source_kind_returns_400() -> None:
    ids = asyncio.run(_create_source(kind="CONFLUENCE"))

    client = TestClient(app)
    response = client.post(f"/sources/{ids['source_id']}/parse")
    _dispose_app_engine()

    assert response.status_code == 400
    assert asyncio.run(_fetch_source_status(ids["source_id"])) == "QUEUED"

    asyncio.run(_cleanup(ids))


def test_missing_source_returns_404() -> None:
    client = TestClient(app)
    response = client.post("/sources/does-not-exist/parse")
    _dispose_app_engine()
    assert response.status_code == 404


def test_parse_failure_stores_actionable_error_on_source() -> None:
    ids = asyncio.run(_create_source())

    async def failing_downloader(_storage_key: str) -> bytes:
        raise BlobDownloadError("simulated blob outage")

    app.dependency_overrides[get_blob_downloader] = lambda: failing_downloader
    client = TestClient(app)
    try:
        client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    async def fetch_parse_error() -> str | None:
        engine = create_async_engine(settings.database_url)
        try:
            async with AsyncSession(engine) as session:
                source = await session.get(Source, ids["source_id"])
                assert source is not None
                return source.parseError
        finally:
            await engine.dispose()

    assert "simulated blob outage" in (asyncio.run(fetch_parse_error()) or "")
    asyncio.run(_cleanup(ids))


def test_duplicate_fragments_within_a_source_are_deduped_on_parse() -> None:
    ids = asyncio.run(_create_source())

    def docx_with_repeated_boilerplate() -> bytes:
        document = Document()
        document.add_paragraph("Company Confidential")
        document.add_paragraph("Real requirement one.")
        document.add_paragraph("Company Confidential")
        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    async def fake_downloader(_storage_key: str) -> bytes:
        return docx_with_repeated_boilerplate()

    app.dependency_overrides[get_blob_downloader] = lambda: fake_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 200
    assert response.json()["chunk_count"] == 2  # duplicate boilerplate dropped

    asyncio.run(_cleanup(ids))


def test_ingestion_summary_reports_sources_and_fragment_counts() -> None:
    ids = asyncio.run(_create_source())

    async def fake_downloader(_storage_key: str) -> bytes:
        return _sample_docx_bytes()

    app.dependency_overrides[get_blob_downloader] = lambda: fake_downloader
    client = TestClient(app)
    try:
        client.post(f"/sources/{ids['source_id']}/parse")
        # Each TestClient call runs on its own event loop — dispose the app engine's
        # pooled (loop-bound) connections between calls, same as between tests.
        _dispose_app_engine()
        summary = client.get(f"/projects/{ids['project_id']}/ingestion-summary")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert summary.status_code == 200
    payload = summary.json()
    assert payload["source_count"] == 1
    assert payload["by_status"] == {"PARSED": 1}
    assert payload["fragment_count"] == 1
    assert payload["sources"][0]["name"] == "test.docx"
    assert payload["last_updated"] is not None

    asyncio.run(_cleanup(ids))


def test_reparse_project_reprocesses_all_parseable_sources() -> None:
    ids = asyncio.run(_create_source())

    async def fake_downloader(_storage_key: str) -> bytes:
        return _sample_docx_bytes()

    app.dependency_overrides[get_blob_downloader] = lambda: fake_downloader
    client = TestClient(app)
    try:
        response = client.post(f"/projects/{ids['project_id']}/reparse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert response.status_code == 200
    results = response.json()["results"]
    assert len(results) == 1
    assert results[0]["status"] == "PARSED"
    assert results[0]["chunk_count"] == 1

    asyncio.run(_cleanup(ids))


def test_parsing_a_new_version_computes_and_persists_a_diff() -> None:
    ids = asyncio.run(_create_source())

    async def v1_downloader(_storage_key: str) -> bytes:
        return _docx_bytes(["Unchanged paragraph.", "Old scope text."])

    app.dependency_overrides[get_blob_downloader] = lambda: v1_downloader
    client = TestClient(app)
    try:
        client.post(f"/sources/{ids['source_id']}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    v2_id = asyncio.run(
        _create_new_version(ids["project_id"], ids["source_id"], previous_version=1, storage_key="v2/key.docx")
    )

    async def v2_downloader(_storage_key: str) -> bytes:
        return _docx_bytes(["Unchanged paragraph.", "New scope text.", "Brand new paragraph."])

    app.dependency_overrides[get_blob_downloader] = lambda: v2_downloader
    client = TestClient(app)
    try:
        parse_response = client.post(f"/sources/{v2_id}/parse")
    finally:
        app.dependency_overrides.pop(get_blob_downloader, None)
        _dispose_app_engine()

    assert parse_response.status_code == 200

    client = TestClient(app)
    diff_response = client.get(f"/sources/{v2_id}/diff")
    _dispose_app_engine()

    assert diff_response.status_code == 200
    body = diff_response.json()
    assert body["source_id"] == v2_id
    assert body["previous_source_id"] == ids["source_id"]
    assert body["unchanged_count"] == 1
    assert body["modified_count"] == 1
    assert body["added_count"] == 1
    assert body["removed_count"] == 0

    # Previous version's RawRequirements remain queryable, untouched.
    assert asyncio.run(_count_raw_requirements(ids["source_id"])) == 2
    assert asyncio.run(_count_raw_requirements(v2_id)) == 3

    asyncio.run(_cleanup(ids, extra_source_ids=[v2_id]))


def test_diff_returns_404_for_a_source_with_no_previous_version() -> None:
    ids = asyncio.run(_create_source())

    client = TestClient(app)
    response = client.get(f"/sources/{ids['source_id']}/diff")
    _dispose_app_engine()

    assert response.status_code == 404

    asyncio.run(_cleanup(ids))
