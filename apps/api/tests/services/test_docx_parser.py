import io

import pytest
from docx import Document

from app.services.parsing.docx_parser import DocxParseError, parse_docx


def _build_docx(build_fn: object) -> bytes:
    document = Document()
    build_fn(document)  # type: ignore[operator]
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_multi_section_doc_produces_correct_nested_section_paths() -> None:
    def build(document: Document) -> None:
        document.add_heading("1. Overview", level=1)
        document.add_paragraph("Overview text.")
        document.add_heading("1.1 Scope", level=2)
        document.add_paragraph("Scope text.")
        document.add_heading("2. Requirements", level=1)
        document.add_paragraph("Requirement text.")

    chunks = parse_docx(_build_docx(build))

    assert [c.section_path for c in chunks] == [
        "1. Overview",
        "1. Overview/1.1 Scope",
        "2. Requirements",
    ]
    assert [c.text for c in chunks] == ["Overview text.", "Scope text.", "Requirement text."]


def test_order_is_ascending_and_reflects_document_order() -> None:
    def build(document: Document) -> None:
        document.add_paragraph("First.")
        document.add_paragraph("Second.")
        document.add_paragraph("Third.")

    chunks = parse_docx(_build_docx(build))
    assert [c.order for c in chunks] == [0, 1, 2]
    assert [c.text for c in chunks] == ["First.", "Second.", "Third."]


def test_table_is_extracted_as_structured_text_not_dropped() -> None:
    def build(document: Document) -> None:
        document.add_paragraph("Before table.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Feature"
        table.cell(0, 1).text = "Priority"
        table.cell(1, 0).text = "Login"
        table.cell(1, 1).text = "High"
        document.add_paragraph("After table.")

    chunks = parse_docx(_build_docx(build))
    table_chunk = chunks[1]
    assert "Feature | Priority" in table_chunk.text
    assert "Login | High" in table_chunk.text
    assert [c.text for c in chunks] == ["Before table.", table_chunk.text, "After table."]


def test_header_and_footer_text_is_extracted() -> None:
    def build(document: Document) -> None:
        document.sections[0].header.paragraphs[0].text = "Company Confidential"
        document.sections[0].footer.paragraphs[0].text = "Page footer text"
        document.add_paragraph("Body text.")

    chunks = parse_docx(_build_docx(build))
    section_paths = {c.section_path for c in chunks}
    assert "#header" in section_paths
    assert "#footer" in section_paths


def test_empty_paragraphs_are_skipped() -> None:
    def build(document: Document) -> None:
        document.add_paragraph("Real text.")
        document.add_paragraph("")
        document.add_paragraph("   ")
        document.add_paragraph("More text.")

    chunks = parse_docx(_build_docx(build))
    assert [c.text for c in chunks] == ["Real text.", "More text."]


def test_falls_back_to_paragraph_index_when_no_heading_context() -> None:
    def build(document: Document) -> None:
        document.add_paragraph("No heading above this.")

    chunks = parse_docx(_build_docx(build))
    assert chunks[0].section_path == "p.1"


def test_corrupt_file_raises_docx_parse_error_not_silent_empty_result() -> None:
    with pytest.raises(DocxParseError):
        parse_docx(b"this is not a valid docx file")


def test_document_with_no_extractable_content_raises_docx_parse_error() -> None:
    def build(document: Document) -> None:
        pass  # a brand-new Document() has no body paragraphs beyond the empty default

    with pytest.raises(DocxParseError):
        parse_docx(_build_docx(build))
